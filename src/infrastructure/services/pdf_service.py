import os
import tempfile
from docx import Document
from docx2pdf import convert
from openpyxl import load_workbook
import platform
from typing import Optional

from src.core.constants import TEMPLATE_PATH, TEMPLATE_RESPALDO_PATH, TEMPLATE_DEVOLUCION_PATH, MESES
from src.domain.models import TransactionRecord
from src.infrastructure.services.file_system import FileSystemService

class DocumentService:
    @staticmethod
    def _set_cell_text(cell, text: str):
        """Escribe texto en una celda de tabla sin perder el formato de párrafo existente."""
        if not cell.paragraphs:
            cell.add_paragraph()
            
        paragraph = cell.paragraphs[0]
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    @staticmethod
    def __get_field(attr, data, deduplicate=False):
        vals = [getattr(d, attr) for d in data.devices if getattr(d, attr)]
        if not vals:
            return ""
        if deduplicate:
            unique_vals = []
            for v in vals:
                if v not in unique_vals:
                    unique_vals.append(v)
            return "\n".join(unique_vals)
        else:
            return "\n".join(vals)

    @staticmethod
    def generate_acta_entrega(data: TransactionRecord) -> Optional[str]:
        """Genera un PDF rellenando la plantilla DOCX 'ACTA ENTREGA.docx'."""
        try:
            doc = Document(TEMPLATE_PATH)

            # --- Párrafo 1: Declaración ---
            p1 = doc.paragraphs[1]
            p1.runs[2].text = data.person.name + " "
            p1.runs[5].text = " " + data.person.ci_document + " "

            # --- Párrafo 15: Fecha ---
            partes_fecha = data.date.split("/")
            if len(partes_fecha) == 3:
                dia, mes_num, anio = partes_fecha
                p15 = doc.paragraphs[15]
                p15.runs[1].text = dia
                p15.runs[3].text = MESES.get(int(mes_num), "___")
                p15.runs[5].text = anio

            # --- Extract device info ---
            tipos = DocumentService.__get_field('device_type', data, deduplicate=True)
            marcas = DocumentService.__get_field('brand', data, deduplicate=True)
            modelos = DocumentService.__get_field('model', data, deduplicate=False)
            
            codigos_list = []
            for d in data.devices:
                c = d.udla_code
                s = d.serial_number
                if c and s:
                    codigos_list.append(f"{c} (S/N: {s})")
                elif c:
                    codigos_list.append(c)
                elif s:
                    codigos_list.append(f"S/N: {s}")
                else:
                    codigos_list.append("-")
            codigos_str = "\n".join(codigos_list)

            # --- Tabla 0: Bienes entregados ---
            t0 = doc.tables[0]
            DocumentService._set_cell_text(t0.rows[0].cells[1], tipos)
            DocumentService._set_cell_text(t0.rows[1].cells[1], marcas)
            DocumentService._set_cell_text(t0.rows[2].cells[1], modelos)
            DocumentService._set_cell_text(t0.rows[3].cells[1], codigos_str)

            DocumentService._set_cell_text(t0.rows[0].cells[2], "")
            DocumentService._set_cell_text(t0.rows[1].cells[2], "")
            DocumentService._set_cell_text(t0.rows[2].cells[2], "")
            DocumentService._set_cell_text(t0.rows[3].cells[2], "")

            DocumentService._set_cell_text(t0.rows[4].cells[1], data.ticket.ticket)
            DocumentService._set_cell_text(t0.rows[5].cells[1], data.date)
            DocumentService._set_cell_text(t0.rows[6].cells[1], data.ticket.observations)

            # --- Tabla 1: Firma ---
            t1 = doc.tables[1]
            DocumentService._set_cell_text(t1.rows[1].cells[1], data.person.name)
            DocumentService._set_cell_text(t1.rows[2].cells[1], data.person.ci_document)

            # --- Guardar ---
            nombre_archivo = "".join(c for c in f"ACTA ENTREGA {data.person.name}" if c not in r'\/:*?"<>|')
            route = FileSystemService.get_records_path("ACTAS DE ENTREGA")
            if not route:
                raise Exception("Could not create output directory")

            temp_dir = tempfile.gettempdir()
            temp_docx = os.path.join(temp_dir, f"{nombre_archivo}.docx")
            output_pdf = os.path.join(route, f"{nombre_archivo}.pdf")

            doc.save(temp_docx)
            convert(temp_docx, output_pdf)

            if os.path.exists(temp_docx):
                os.remove(temp_docx)

            return output_pdf
        except Exception as e:
            print(f"Error generating acta de entrega: {e}")
            raise

    @staticmethod
    def generate_acta_respaldo(data: TransactionRecord) -> Optional[str]:
        """Genera un PDF rellenando la plantilla de respaldo."""
        try:
            doc = Document(TEMPLATE_RESPALDO_PATH)

            partes_fecha = data.date.split("/")
            if len(partes_fecha) == 3:
                dia, mes, anio = partes_fecha
                p4 = doc.paragraphs[4]
                p4.runs[3].text = dia
                p4.runs[5].text = mes
                p4.runs[7].text = anio
                p4.runs[8].text = ""

            p7 = doc.paragraphs[7]
            p7.runs[3].text = data.person.name + " "

            p11 = doc.paragraphs[11]
            p11.runs[2].text = data.person.name
            p11.runs[4].text = data.reason or "realizar un cambio y/o formateo del equipo"
            p11.runs[5].text = ""
            p11.runs[6].text = "."

            p18 = doc.paragraphs[18]
            p18.runs[2].text = data.person.name

            nombre_archivo = "".join(c for c in f"ACTA RESPALDO {data.person.name}" if c not in r'\/:*?"<>|')
            route = FileSystemService.get_records_path("ACTAS DE RESPALDO")
            if not route:
                raise Exception("Could not create output directory")

            temp_dir = tempfile.gettempdir()
            temp_docx = os.path.join(temp_dir, f"{nombre_archivo}.docx")
            output_pdf = os.path.join(route, f"{nombre_archivo}.pdf")

            doc.save(temp_docx)
            convert(temp_docx, output_pdf)

            if os.path.exists(temp_docx):
                os.remove(temp_docx)

            return output_pdf
        except Exception as e:
            print(f"Error generating acta de respaldo: {e}")
            raise

    @staticmethod
    def generate_acta_devolucion(data: TransactionRecord) -> Optional[str]:
        """Genera el Excel de devolución y lo devuelve (la ruta)."""
        try:
            nombre_archivo = "".join(c for c in f"ACTA DEVOLUCION {data.person.name}" if c not in r'\/:*?"<>|')
            route = FileSystemService.get_records_path("ACTAS DE DEVOLUCION")
            if not route:
                raise Exception("Could not create output directory")
                
            wb = load_workbook(TEMPLATE_DEVOLUCION_PATH)
            # Add mapping logic if the template has specific cells to fill
            # Currently it just saves a copy based on the original code
            
            output_xlsx = os.path.join(route, f"{nombre_archivo}.xlsx")
            wb.save(output_xlsx)
            return output_xlsx
        except Exception as e:
            print(f"Error generating acta de devolucion: {e}")
            raise
